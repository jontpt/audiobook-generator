"""
services/text_extraction.py
Handles extraction of raw text from PDF, DOCX, ePub, and plain-text files.
Returns a tuple:
  (list[dict], dict[str, str], dict[str, str])
  ├─ list[dict]   → chapters: [{"title": str, "paragraphs": [str]}, ...]
  ├─ dict[str,str]→ char_declarations: {"CharName": "male"|"female"|"neutral", ...}
  │                 (empty dict when no CHARACTERS: block is present)
  └─ dict[str,str]→ voice_hints: {"CharName": "voice name/id hint", ...}
                    (empty dict when no VOICE line is present)

CHARACTERS block format (place anywhere at top of file, before the story):
──────────────────────────────────────────────
CHARACTERS:
Spade: male
Wonderly: female
Effie: female
Archer: male
END CHARACTERS
──────────────────────────────────────────────
"""
from __future__ import annotations
import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# ── CHARACTERS block parser ───────────────────────────────────────────────────
# Matches the entire CHARACTERS: ... END CHARACTERS block (case-insensitive,
# with or without the END CHARACTERS terminator — falls back to blank-line end).

_CHAR_BLOCK_RE = re.compile(
    r'CHARACTERS\s*:\s*\n'           # opening line
    r'(.*?)'                          # content (non-greedy)
    r'(?:END\s+CHARACTERS\s*(?:\n|$)|(?:\n{2,}|\Z))',  # closing: explicit end or blank line
    re.IGNORECASE | re.DOTALL
)
_CHAR_LINE_RE = re.compile(
    r'^\s*([A-Za-z][A-Za-z\s\'\-\.]+?)\s*:\s*(male|female|neutral|m|f|n)\s*$',
    re.IGNORECASE
)
_VOICE_LINE_RE = re.compile(
    r'^\s*([A-Za-z][A-Za-z\s\'\-\.]+?)\s*:\s*(?:voice|voice_id)\s*=\s*(.+?)\s*$',
    re.IGNORECASE
)
_GENDER_ALIASES: dict[str, str] = {
    "m": "male", "f": "female", "n": "neutral",
    "male": "male", "female": "female", "neutral": "neutral",
}


def _parse_char_declarations(raw_text: str) -> tuple[dict[str, str], dict[str, str], str]:
    """
    Extract the CHARACTERS: block from raw_text.
    Returns (declarations_dict, voice_hints_dict, cleaned_text_without_block).
    declarations_dict maps CharName → "male"|"female"|"neutral".
    voice_hints_dict maps CharName → requested voice label/id string.
    If no block is found returns ({}, {}, raw_text unchanged).
    """
    m = _CHAR_BLOCK_RE.search(raw_text)
    if not m:
        return {}, {}, raw_text

    block_content = m.group(1)
    declarations: dict[str, str] = {}
    voice_hints: dict[str, str] = {}
    for line in block_content.splitlines():
        lm = _CHAR_LINE_RE.match(line)
        if lm:
            name   = lm.group(1).strip()
            gender = _GENDER_ALIASES.get(lm.group(2).strip().lower(), "neutral")
            declarations[name] = gender
            logger.debug(f"CHARACTERS block: {name!r} → {gender}")
            continue
        vm = _VOICE_LINE_RE.match(line)
        if vm:
            name = vm.group(1).strip()
            voice_hints[name] = vm.group(2).strip()
            logger.debug(f"CHARACTERS block voice hint: {name!r} → {voice_hints[name]!r}")

    # Remove the entire block from the text so it doesn't pollute chapters
    cleaned = raw_text[:m.start()] + raw_text[m.end():]
    if declarations:
        logger.info(f"Parsed CHARACTERS block: {list(declarations.keys())}")
    if voice_hints:
        logger.info(f"Parsed CHARACTERS voice hints: {list(voice_hints.keys())}")
    return declarations, voice_hints, cleaned


# ── Chapter detection patterns ───────────────────────────────────────────────
CHAPTER_PATTERNS = [
    re.compile(r'^chapter\s+\w+', re.IGNORECASE),
    re.compile(r'^(part|book|volume|act|scene)\s+\w+', re.IGNORECASE),
    re.compile(r'^\d+\.\s+\w+'),          # "1. Title"
    re.compile(r'^[IVXLCDM]+\.\s+\w+'),   # Roman numerals
]

HEADER_FOOTER_PATTERNS = [
    re.compile(r'^\d+$'),                  # Lone page number
    re.compile(r'^page\s+\d+', re.I),
    re.compile(r'^[-–—]{3,}$'),            # Dividers
]


def _is_chapter_heading(line: str) -> bool:
    stripped = line.strip()
    return any(p.match(stripped) for p in CHAPTER_PATTERNS)


def _is_noise(line: str) -> bool:
    stripped = line.strip()
    return any(p.match(stripped) for p in HEADER_FOOTER_PATTERNS)


def _split_into_chapters(raw_text: str) -> list[dict]:
    """
    Split raw text into chapters.
    Each chapter: {"title": str, "paragraphs": [str]}
    """
    lines = raw_text.split('\n')
    chapters: list[dict] = []
    current_title = "Chapter 1"
    current_paras: list[str] = []
    buffer: list[str] = []

    def flush_buffer():
        block = ' '.join(buffer).strip()
        if block:
            current_paras.append(block)
        buffer.clear()

    for line in lines:
        clean = line.strip()
        if not clean or _is_noise(clean):
            flush_buffer()
            continue

        if _is_chapter_heading(clean):
            flush_buffer()
            if current_paras:
                chapters.append({"title": current_title, "paragraphs": current_paras.copy()})
            current_title = clean
            current_paras = []
        else:
            buffer.append(clean)

    flush_buffer()
    if current_paras:
        chapters.append({"title": current_title, "paragraphs": current_paras})

    if not chapters:
        # No chapter markers found — treat whole text as one chapter
        all_paras = [p.strip() for p in raw_text.split('\n\n') if p.strip()]
        chapters = [{"title": "Chapter 1", "paragraphs": all_paras}]

    return chapters


# ── PDF ──────────────────────────────────────────────────────────────────────

def extract_from_pdf(file_path: Path) -> tuple[list[dict], dict[str, str], dict[str, str]]:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(file_path))
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text("text"))
        doc.close()
        raw = '\n'.join(pages_text)
        char_decls, voice_hints, cleaned = _parse_char_declarations(raw)
        return _split_into_chapters(cleaned), char_decls, voice_hints
    except ImportError:
        logger.warning("PyMuPDF not installed, falling back to basic PDF reading")
        return _fallback_pdf(file_path)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise


def _fallback_pdf(file_path: Path) -> tuple[list[dict], dict[str, str], dict[str, str]]:
    """Minimal fallback using pdfminer if available."""
    try:
        from pdfminer.high_level import extract_text as pm_extract
        raw = pm_extract(str(file_path))
        char_decls, voice_hints, cleaned = _parse_char_declarations(raw)
        return _split_into_chapters(cleaned), char_decls, voice_hints
    except Exception as e:
        raise RuntimeError(f"Cannot extract PDF text: {e}")


# ── DOCX ─────────────────────────────────────────────────────────────────────

def extract_from_docx(file_path: Path) -> tuple[list[dict], dict[str, str], dict[str, str]]:
    try:
        from docx import Document
        doc = Document(str(file_path))
        # First pass: collect all paragraph text and check for CHARACTERS block
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        char_decls, voice_hints, _ = _parse_char_declarations(all_text)

        # Build a set of lines that are inside the CHARACTERS block (to skip)
        char_block_lines: set[str] = set()
        m = _CHAR_BLOCK_RE.search(all_text)
        if m:
            char_block_lines = {ln.strip() for ln in m.group(0).splitlines() if ln.strip()}

        chapters: list[dict] = []
        current_title = "Chapter 1"
        current_paras: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Skip lines belonging to the CHARACTERS block
            if text in char_block_lines:
                continue
            # Use heading styles as chapter breaks
            if para.style.name.startswith('Heading 1') or _is_chapter_heading(text):
                if current_paras:
                    chapters.append({"title": current_title, "paragraphs": current_paras.copy()})
                current_title = text
                current_paras = []
            else:
                current_paras.append(text)

        if current_paras:
            chapters.append({"title": current_title, "paragraphs": current_paras})

        if not chapters:
            chapters = _split_into_chapters(all_text)

        return chapters, char_decls, voice_hints
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise


# ── ePub ─────────────────────────────────────────────────────────────────────

def extract_from_epub(file_path: Path) -> tuple[list[dict], dict[str, str], dict[str, str]]:
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(str(file_path))
        chapters: list[dict] = []
        all_char_decls: dict[str, str] = {}
        all_voice_hints: dict[str, str] = {}

        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            # Try to get chapter title from h1/h2/h3
            heading = soup.find(['h1', 'h2', 'h3'])
            title = heading.get_text().strip() if heading else f"Chapter {len(chapters)+1}"
            # Extract paragraphs
            raw_paras = [p.get_text().strip() for p in soup.find_all('p') if p.get_text().strip()]
            # Parse CHARACTERS block from this section's text
            section_text = '\n'.join(raw_paras)
            char_decls, voice_hints, cleaned_text = _parse_char_declarations(section_text)
            all_char_decls.update(char_decls)
            all_voice_hints.update(voice_hints)
            paras = [p.strip() for p in cleaned_text.split('\n') if p.strip()]
            if paras:
                chapters.append({"title": title, "paragraphs": paras})

        return chapters, all_char_decls, all_voice_hints
    except Exception as e:
        logger.error(f"ePub extraction error: {e}")
        raise


# ── Plain text ────────────────────────────────────────────────────────────────

def extract_from_txt(file_path: Path) -> tuple[list[dict], dict[str, str], dict[str, str]]:
    raw = file_path.read_text(encoding='utf-8', errors='replace')
    char_decls, voice_hints, cleaned = _parse_char_declarations(raw)
    return _split_into_chapters(cleaned), char_decls, voice_hints


# ── Dispatcher ───────────────────────────────────────────────────────────────

def extract_text(file_path: Path) -> tuple[list[dict], dict[str, str], dict[str, str]]:
    """
    Main entry point. Detect file type and extract chapters + character declarations.

    Returns:
        (chapters, char_declarations, voice_hints)
        chapters:          list of {"title": str, "paragraphs": [str]}
        char_declarations: {"CharName": "male"|"female"|"neutral"} (may be empty)
        voice_hints:       {"CharName": "voice hint"} (may be empty)
    """
    ext = file_path.suffix.lower()
    extractors = {
        '.pdf':  extract_from_pdf,
        '.docx': extract_from_docx,
        '.epub': extract_from_epub,
        '.txt':  extract_from_txt,
    }
    extractor = extractors.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file type: {ext}")

    logger.info(f"Extracting text from {file_path.name} ({ext})")
    chapters, char_declarations, voice_hints = extractor(file_path)
    logger.info(
        f"Extracted {len(chapters)} chapters"
        + (f", {len(char_declarations)} declared characters" if char_declarations else "")
        + (f", {len(voice_hints)} voice hints" if voice_hints else "")
    )
    return chapters, char_declarations, voice_hints

